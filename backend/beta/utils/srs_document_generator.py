"""
SRS Document Generator Utility with Table of Contents

This module provides functionality to generate professionally formatted 
Software Requirements Specification (SRS) documents from JSON data with
automatic Table of Contents and page numbering.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
from typing import Dict, Any, List, Optional


class SRSDocumentGenerator:
    """Generate SRS documents from JSON data with proper formatting and TOC."""
    
    def __init__(self, project_name: str, authors: List[str] = None, organization: str = "Organization Name"):
        """
        Initialize the SRS document generator.
        
        Args:
            project_name: Name of the project for headers
            authors: List of document author names (default: ["Author Name"])
            organization: Organization name (default: "Organization Name")
        """
        self.project_name = str(project_name or "Project")
        self.authors = [str(a) for a in (authors or ["Author Name"])]
        self.organization = str(organization or "Organization Name")
        self.doc = Document()
        self.image_paths = {}  # set before adding sections; may include system_context, system_architecture, use_case, user_workflow, security_flow, data_erd
        self._setup_document()
        self._setup_styles()

    def _apply_section_layout(self, section):
        """Apply consistent page geometry to a single section."""
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.header_distance = Inches(0.5)
        section.footer_distance = Inches(0.5)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

    def _apply_layout_to_all_sections(self):
        """Ensure every section (all pages) uses consistent layout."""
        for section in self.doc.sections:
            self._apply_section_layout(section)
        
    def _setup_document(self):
        """Configure document-level settings (margins, page size)."""
        sections = self.doc.sections
        for section in sections:
            self._apply_section_layout(section)
            
    def _setup_styles(self):
        """Set up custom styles for the document."""
        styles = self.doc.styles
        
        # Modify Normal style
        normal_style = styles['Normal']
        normal_font = normal_style.font
        normal_font.name = 'Arial'
        normal_font.size = Pt(11)
        normal_style.paragraph_format.line_spacing = 1.25
        normal_style.paragraph_format.space_after = Pt(8)
        normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Heading 1 style
        try:
            h1_style = styles['Heading 1']
        except KeyError:
            h1_style = styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
        h1_font = h1_style.font
        h1_font.name = 'Arial'
        h1_font.size = Pt(16)
        h1_font.bold = True
        h1_font.color.rgb = RGBColor(15, 63, 102)
        h1_style.paragraph_format.space_before = Pt(24)
        h1_style.paragraph_format.space_after = Pt(12)
        h1_style.paragraph_format.keep_with_next = True
        
        # Heading 2 style
        try:
            h2_style = styles['Heading 2']
        except KeyError:
            h2_style = styles.add_style('Heading 2', WD_STYLE_TYPE.PARAGRAPH)
        h2_font = h2_style.font
        h2_font.name = 'Arial'
        h2_font.size = Pt(13)
        h2_font.bold = True
        h2_font.color.rgb = RGBColor(20, 20, 20)
        h2_style.paragraph_format.space_before = Pt(14)
        h2_style.paragraph_format.space_after = Pt(9)
        h2_style.paragraph_format.keep_with_next = True
        
        # Heading 3 style
        try:
            h3_style = styles['Heading 3']
        except KeyError:
            h3_style = styles.add_style('Heading 3', WD_STYLE_TYPE.PARAGRAPH)
        h3_font = h3_style.font
        h3_font.name = 'Arial'
        h3_font.size = Pt(12)
        h3_font.bold = True
        h3_font.color.rgb = RGBColor(0, 0, 0)
        h3_style.paragraph_format.space_before = Pt(12)
        h3_style.paragraph_format.space_after = Pt(6)
        h3_style.paragraph_format.keep_with_next = True
        
        # TOC Heading style
        try:
            toc_heading_style = styles['TOC Heading']
        except KeyError:
            toc_heading_style = styles.add_style('TOC Heading', WD_STYLE_TYPE.PARAGRAPH)
        toc_heading_font = toc_heading_style.font
        toc_heading_font.name = 'Arial'
        toc_heading_font.size = Pt(15)
        toc_heading_font.bold = True
        toc_heading_font.color.rgb = RGBColor(15, 63, 102)
        toc_heading_style.paragraph_format.space_before = Pt(24)
        toc_heading_style.paragraph_format.space_after = Pt(12)
        toc_heading_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Figure caption style
        try:
            caption_style = styles['Caption']
        except KeyError:
            caption_style = styles.add_style('Caption', WD_STYLE_TYPE.PARAGRAPH)
        caption_font = caption_style.font
        caption_font.name = 'Arial'
        caption_font.size = Pt(10)
        caption_font.italic = True
        caption_font.color.rgb = RGBColor(80, 80, 80)
        caption_style.paragraph_format.space_before = Pt(4)
        caption_style.paragraph_format.space_after = Pt(10)
        caption_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # List bullet style cleanup for consistent indentation
        try:
            bullet_style = styles['List Bullet']
            bullet_style.paragraph_format.left_indent = Inches(0.25)
            bullet_style.paragraph_format.first_line_indent = Inches(-0.15)
            bullet_style.paragraph_format.space_after = Pt(4)
        except KeyError:
            pass

    def _add_figure(self, path: Path, caption: str, width: float = 5.8):
        """Insert a centered figure with a consistent caption style."""
        self.doc.add_picture(str(path), width=Inches(width))
        image_para = self.doc.paragraphs[-1]
        image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        caption_para = self.doc.add_paragraph()
        caption_para.style = 'Caption'
        run = caption_para.add_run(caption)
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        run.italic = True
        
    def _add_visual_grid(self, items, columns: int = 2, image_width: float = 2.8):
        """Insert a compact grid of figures with captions."""
        if not items:
            return
        rows = (len(items) + columns - 1) // columns
        table = self.doc.add_table(rows=rows, cols=columns)
        table.autofit = True
        idx = 0
        for r in range(rows):
            for c in range(columns):
                cell = table.cell(r, c)
                cell.text = ""
                if idx >= len(items):
                    continue
                path, caption = items[idx]
                para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if path and Path(path).exists():
                    run = para.add_run()
                    run.add_picture(str(path), width=Inches(image_width))
                    cap = cell.add_paragraph(caption)
                    cap.style = "Caption"
                else:
                    placeholder = cell.add_paragraph(caption)
                    placeholder.style = "Caption"
                idx += 1

    def _add_header_footer(self):
        """Add header and footer to all pages including title and TOC."""
        # Get all sections
        sections = self.doc.sections
        
        # Configure all sections (including title and TOC)
        for idx, section in enumerate(sections):
            # Keep front matter (title + TOC) clean
            if idx < 2:
                continue

            # Header - add to ALL pages
            header = section.header
            header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            header_para.text = f"{self.project_name} | Software Requirements Specification"
            header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if header_para.runs:
                header_para.runs[0].font.size = Pt(10)
                header_para.runs[0].font.name = 'Arial'
                header_para.runs[0].font.color.rgb = RGBColor(102, 102, 102)
            
            # Footer with page number - add to ALL pages
            footer = section.footer
            footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Clear existing content in a version-safe way.
            # Some python-docx versions do not expose Paragraph.clear().
            if hasattr(footer_para, "clear"):
                footer_para.clear()
            else:
                para_element = footer_para._element
                for child in list(para_element):
                    para_element.remove(child)
            
            # Add "Page X of Y" fields
            footer_para.add_run("Page ").font.name = 'Arial'

            run = footer_para.add_run()
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            
            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = "PAGE"
            
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            
            run._r.append(fldChar1)
            run._r.append(instrText)
            run._r.append(fldChar2)
            run.font.size = Pt(10)
            run.font.name = 'Arial'
            run.font.color.rgb = RGBColor(102, 102, 102)

            footer_para.add_run(" of ").font.name = 'Arial'

            run_total = footer_para.add_run()
            total_fld_begin = OxmlElement('w:fldChar')
            total_fld_begin.set(qn('w:fldCharType'), 'begin')
            total_instr = OxmlElement('w:instrText')
            total_instr.set(qn('xml:space'), 'preserve')
            total_instr.text = "NUMPAGES"
            total_fld_end = OxmlElement('w:fldChar')
            total_fld_end.set(qn('w:fldCharType'), 'end')
            run_total._r.append(total_fld_begin)
            run_total._r.append(total_instr)
            run_total._r.append(total_fld_end)
            run_total.font.size = Pt(10)
            run_total.font.name = 'Arial'
            run_total.font.color.rgb = RGBColor(102, 102, 102)

    def _add_horizontal_rule(self):
        line = self.doc.add_paragraph()
        line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = line.add_run("".ljust(90, "_"))
        run.font.name = 'Arial'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(170, 170, 170)

    def _insert_toc_field(self):
        """Insert an automatic TOC field that Word can update on open."""
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run()

        fld_char_begin = OxmlElement('w:fldChar')
        fld_char_begin.set(qn('w:fldCharType'), 'begin')
        instr_text = OxmlElement('w:instrText')
        instr_text.set(qn('xml:space'), 'preserve')
        instr_text.text = r'TOC \o "1-3" \h \z \u'
        fld_char_sep = OxmlElement('w:fldChar')
        fld_char_sep.set(qn('w:fldCharType'), 'separate')
        text = OxmlElement('w:t')
        text.text = "Right-click and update field to generate Table of Contents."
        fld_char_end = OxmlElement('w:fldChar')
        fld_char_end.set(qn('w:fldCharType'), 'end')

        run._r.append(fld_char_begin)
        run._r.append(instr_text)
        run._r.append(fld_char_sep)
        run._r.append(text)
        run._r.append(fld_char_end)
        
    def _add_title_page(self):
        """Add title page for the SRS document."""
        from datetime import datetime

        self.doc.add_paragraph()
        self._add_horizontal_rule()

        # Document type
        doc_type = self.doc.add_paragraph()
        doc_type.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = doc_type.add_run("SOFTWARE REQUIREMENTS SPECIFICATION")
        run.font.name = 'Arial'
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.color.rgb = RGBColor(15, 63, 102)
        doc_type.paragraph_format.space_before = Pt(72)
        doc_type.paragraph_format.space_after = Pt(24)

        # Subtitle
        subtitle = self.doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run("Prepared in IEEE 830 style")
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        run.font.italic = True
        run.font.color.rgb = RGBColor(96, 96, 96)
        subtitle.paragraph_format.space_after = Pt(28)

        # Title
        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(self.project_name)
        run.font.name = 'Arial'
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(25, 25, 25)
        
        # Add spacing
        title.paragraph_format.space_after = Pt(44)

        meta_heading = self.doc.add_paragraph()
        meta_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = meta_heading.add_run("Document Information")
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(15, 63, 102)
        meta_heading.paragraph_format.space_after = Pt(16)

        prepared_by = self.doc.add_paragraph()
        prepared_by.alignment = WD_ALIGN_PARAGRAPH.CENTER
        authors_text = ", ".join(self.authors)
        run = prepared_by.add_run(f"Prepared by: {authors_text}")
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        
        prepared_by.paragraph_format.space_after = Pt(12)
        
        # Organization
        organization = self.doc.add_paragraph()
        organization.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = organization.add_run(f"Organization: {self.organization}")
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        
        organization.paragraph_format.space_after = Pt(12)

        version = self.doc.add_paragraph()
        version.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = version.add_run("Version: 1.0")
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        version.paragraph_format.space_after = Pt(12)
        
        # Date created
        date_created = self.doc.add_paragraph()
        date_created.alignment = WD_ALIGN_PARAGRAPH.CENTER
        today_date = datetime.now().strftime("%m/%d/%Y")
        run = date_created.add_run(f"Date Created: {today_date}")
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        date_created.paragraph_format.space_after = Pt(26)

        # Document control table for enterprise formatting
        control_table = self.doc.add_table(rows=3, cols=2)
        control_table.style = 'Table Grid'
        control_table.autofit = True
        control_rows = [
            ("Document ID", f"SRS-{self.project_name[:20].upper().replace(' ', '-')}-001"),
            ("Document Status", "Draft"),
            ("Prepared For", self.organization),
        ]
        for i, (k, v) in enumerate(control_rows):
            c0 = control_table.cell(i, 0)
            c1 = control_table.cell(i, 1)
            c0.text = k
            c1.text = v
            if c0.paragraphs and c0.paragraphs[0].runs:
                c0.paragraphs[0].runs[0].bold = True
                c0.paragraphs[0].runs[0].font.name = 'Arial'
                c0.paragraphs[0].runs[0].font.size = Pt(10)
            if c1.paragraphs and c1.paragraphs[0].runs:
                c1.paragraphs[0].runs[0].font.name = 'Arial'
                c1.paragraphs[0].runs[0].font.size = Pt(10)

        self._add_horizontal_rule()
        
        # Add section break (new page) after title page
        self.doc.add_section()
    
    def _add_table_of_contents(self):
        """Add Table of Contents after the title page."""
        toc_heading = self.doc.add_paragraph()
        toc_heading.style = 'TOC Heading'
        run = toc_heading.add_run("Table of Contents")
        run.font.name = 'Arial'
        run.font.size = Pt(16)
        run.font.bold = True
        toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        toc_heading.paragraph_format.space_after = Pt(18)
        self._insert_toc_field()

        note = self.doc.add_paragraph()
        note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        note_run = note.add_run("Open in Microsoft Word and use 'Update Field' to refresh page numbers.")
        note_run.font.italic = True
        note_run.font.size = Pt(9)
        note_run.font.color.rgb = RGBColor(128, 128, 128)
        note.paragraph_format.space_before = Pt(10)

        # Section break (new page) after TOC
        new_section = self.doc.add_section()
        new_section.start_type = 2
        # page_number_start is not available in all python-docx versions.
        # Skip hard assignment to avoid generation crashes.
    
    def _set_update_fields_on_open(self):
        """Set the document to update fields when opened."""
        try:
            # Access document settings
            settings_element = self.doc.settings.element
            
            # Create updateFields element if it doesn't exist
            update_fields = settings_element.find(qn('w:updateFields'))
            if update_fields is None:
                update_fields = OxmlElement('w:updateFields')
                update_fields.set(qn('w:val'), 'true')
                settings_element.append(update_fields)
            else:
                update_fields.set(qn('w:val'), 'true')
        except Exception as e:
            # If settings don't exist, create them
            pass
        
    def add_introduction_section(self, intro_data: Dict[str, Any]):
        """
        Add Introduction section to the document.
        
        Args:
            intro_data: Dictionary containing introduction section data
        """
        # Section title
        self.doc.add_heading(intro_data.get('title', '1. Introduction'), level=1)
        
        # 1.1 Purpose
        purpose = intro_data.get('purpose', {})
        self.doc.add_heading(purpose.get('title', '1.1 Purpose'), level=2)
        self.doc.add_paragraph(purpose.get('description', ''))
        
        # 1.2 Scope of the System
        scope = intro_data.get('project_scope', {})
        self.doc.add_heading('1.2 Scope of the System', level=2)
        scope_desc = scope.get('description') or "The system provides a centralized platform for core business operations with secure access, reporting, and monitoring. Features outside the specified requirements are excluded from this version."
        self.doc.add_paragraph(scope_desc)
        included = scope.get('included', [])
        if included:
            self.doc.add_paragraph("Included in scope:", style='Heading 3')
            for item in included:
                self.doc.add_paragraph(item, style='List Bullet')
        excluded = scope.get('excluded', [])
        if excluded:
            self.doc.add_paragraph("Excluded from scope:", style='Heading 3')
            for item in excluded:
                self.doc.add_paragraph(item, style='List Bullet')
        
        # 1.3 Definitions, Acronyms, and Abbreviations
        self.doc.add_heading('1.3 Definitions, Acronyms, and Abbreviations', level=2)
        for term, defn in [("SRS", "Software Requirements Specification"), ("RBAC", "Role-Based Access Control"), ("CRUD", "Create, Read, Update, Delete"), ("UI", "User Interface"), ("API", "Application Programming Interface")]:
            self.doc.add_paragraph(f"{term} – {defn}", style='List Bullet')
        
        # 1.4 Document Conventions (optional)
        conventions = intro_data.get('document_conventions', {})
        self.doc.add_heading('1.4 Document Conventions', level=2)
        # Always add IEEE 830-1998 style convention
        self.doc.add_paragraph("IEEE 830-1998 style", style='List Bullet')
        # Add any additional conventions if provided
        if conventions.get('conventions'):
            for conv in conventions.get('conventions', []):
                if conv and "IEEE" not in conv:  # Avoid duplicate
                    self.doc.add_paragraph(conv, style='List Bullet')
        
        # 1.5 References (optional)
        references = intro_data.get('references', {})
        if references.get('references'):
            self.doc.add_heading(references.get('title', '1.5 References'), level=2)
            for ref in references.get('references', []):
                ref_id = ref.get('id', '')
                ref_desc = ref.get('description', '')
                self.doc.add_paragraph(f"{ref_id}: {ref_desc}", style='List Bullet')

        # 1.6 Live Prototype (New)
        if intro_data.get('live_link'):
            self.doc.add_heading('1.6 Live Prototype', level=2)
            p = self.doc.add_paragraph("You can interact with the generated prototype of this system here: ")
            self.add_hyperlink(p, "Click here to view Live Project Prototype", intro_data.get('live_link'))

    def add_hyperlink(self, paragraph, text, url):
        """
        Add a hyperlink to a paragraph.
        """
        # This gets access to the document.xml.rels file and gets a new relation id value
        part = paragraph.part
        r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

        # Create the w:hyperlink tag and add needed values
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)

        # Create a w:r element and a new w:rPr element
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')

        # Join all the xml elements together add add the required text to the w:r element
        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)

        # Create a new R element and add the hyperlink to it
        r = paragraph.add_run()
        r._r.append(hyperlink)

        # A workaround for the color and underlining
        r.font.color.rgb = RGBColor(0, 0, 255)
        r.font.underline = True
        return hyperlink

    def add_feasibility_section(self):
        """
        Add Feasibility & Cost Estimation (COCOMO) section.
        """
        self.doc.add_heading('2. Feasibility & Cost Estimation', level=1)
        self.doc.add_paragraph("The following cost and effort estimates are based on the COCOMO II model.")

        table = self.doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Metric'
        hdr_cells[1].text = 'Estimated Value'

        # Mock Data (In real app, calculate this based on function points)
        metrics = [
             ("Estimated Lines of Code (KLOC)", "5.2 KLOC"),
             ("Development Effort", "14.5 Person-Months"),
             ("Development Cost", "$45,000"),
             ("Required Schedule", "6 Months")
        ]
        
        for metric, value in metrics:
            row_cells = table.add_row().cells
            row_cells[0].text = metric
            row_cells[1].text = value
    
    def add_overall_description_section(self, desc_data: Dict[str, Any]):
        """
        Add Overall Description section to the document.
        
        Args:
            desc_data: Dictionary containing overall description section data
        """
        # Section title
        self.doc.add_heading(desc_data.get('title', '2. Overall Description'), level=1)
        
        # 2.1 Product Perspective
        perspective = desc_data.get('product_perspective', {})
        self.doc.add_heading(perspective.get('title', '2.1 Product Perspective'), level=2)
        self.doc.add_paragraph(perspective.get('description', ''))
        path = self.image_paths.get('system_context')
        if path and Path(path).exists():
            self._add_figure(
                Path(path),
                "Figure 1: System Context Diagram",
            )
        else:
            self.doc.add_paragraph("Note: The System Context Diagram illustrating interactions between the system and external entities can be added above.", style='Heading 3')
        
        # 2.2 Product Features
        features = desc_data.get('product_features', {})
        self.doc.add_heading(features.get('title', '2.2 Product Features'), level=2)
        feature_list = features.get('features', [])
        for feature in feature_list:
            self.doc.add_paragraph(feature, style='List Bullet')
        
        # 2.3 User Classes and Characteristics
        user_classes = desc_data.get('user_classes_and_characteristics', {})
        self.doc.add_heading(user_classes.get('title', '2.3 User Classes and Characteristics'), level=2)
        classes = user_classes.get('user_classes', [])
        for user_class in classes:
            user_type = user_class.get('user_class', '')
            self.doc.add_paragraph(user_type)
        
        # 2.4 Operating Environment
        environment = desc_data.get('operating_environment', {})
        self.doc.add_heading(environment.get('title', '2.4 Operating Environment'), level=2)
        env_list = environment.get('environments', [])
        for env in env_list:
            self.doc.add_paragraph(env, style='List Bullet')
        
        # 2.5 Design and Implementation Constraints
        constraints = desc_data.get('design_and_implementation_constraints', {})
        self.doc.add_heading(constraints.get('title', '2.5 Design and Implementation Constraints'), level=2)
        constraint_list = constraints.get('constraints', [])
        for constraint in constraint_list:
            self.doc.add_paragraph(constraint, style='List Bullet')
        
        # 2.6 User Documentation
        documentation = desc_data.get('user_documentation', {})
        self.doc.add_heading(documentation.get('title', '2.6 User Documentation'), level=2)
        doc_list = documentation.get('documents', [])
        for doc in doc_list:
            self.doc.add_paragraph(doc, style='List Bullet')
        
        # 2.7 Assumptions and Dependencies
        assumptions = desc_data.get('assumptions_and_dependencies', {})
        self.doc.add_heading(assumptions.get('title', '2.7 Assumptions and Dependencies'), level=2)
        
        assumption_list = assumptions.get('assumptions', [])
        if assumption_list:
            self.doc.add_paragraph("Assumptions:", style='Heading 3')
            for assumption in assumption_list:
                self.doc.add_paragraph(assumption, style='List Bullet')
        
        dependency_list = assumptions.get('dependencies', [])
        if dependency_list:
            self.doc.add_paragraph("Dependencies:", style='Heading 3')
            for dependency in dependency_list:
                self.doc.add_paragraph(dependency, style='List Bullet')
    
    def add_system_architecture_section(self, desc_data: Dict[str, Any] = None):
        """Section 3: System Architecture with architecture diagram."""
        self.doc.add_heading('3. System Architecture', level=1)
        self.doc.add_paragraph(
            "The system follows a layered architecture consisting of: Presentation Layer (Web UI), "
            "Application Layer (Backend services), Data Layer (Database), and External Integration Layer (APIs, third-party services)."
        )
        path = self.image_paths.get('system_architecture')
        if path and Path(path).exists():
            self._add_figure(
                Path(path),
                "Figure 2: System Architecture Diagram",
            )
        else:
            self.doc.add_paragraph("Note: The System Architecture Diagram can be included above.", style='Heading 3')
    
    def add_system_features_section(self, features_data: Dict[str, Any]):
        """
        Add Functional Requirements section (4) with Use Case diagram.
        """
        self.doc.add_heading('4. Functional Requirements', level=1)
        path = self.image_paths.get('use_case')
        if path and Path(path).exists():
            self._add_figure(
                Path(path),
                "Figure 3: Use Case Diagram",
            )
        features = features_data.get('features', [])
        for idx, feature in enumerate(features, 1):
            feature_name = feature.get('feature_name', f'Feature {idx}')
            self.doc.add_heading(f"4.{idx} {feature_name}", level=2)
            
            # Description (if provided)
            description = feature.get('description', '')
            if description:
                self.doc.add_paragraph(f"Description: {description}")
            
            # Simplified functional requirements format: "Support: {feature_name}"
            self.doc.add_paragraph("Functional Requirements:", style='Heading 3')
            self.doc.add_paragraph(f"Support: {feature_name}", style='List Bullet')
    
    def add_user_workflow_section(self):
        """Section 5: User Workflow with workflow diagram."""
        self.doc.add_heading('5. User Workflow', level=1)
        self.doc.add_paragraph(
            "The typical user workflow includes: User logs into the system; System validates credentials; "
            "User accesses role-based dashboard; User performs permitted actions; System processes and stores data; User views results or reports."
        )
        path = self.image_paths.get('user_workflow')
        if path and Path(path).exists():
            self._add_figure(
                Path(path),
                "Figure 4: User Workflow Diagram",
            )
        else:
            self.doc.add_paragraph("Note: The User Flow / Workflow Diagram can be included above.", style='Heading 3')

    def add_visual_overview_section(self):
        """A compact, diagram-heavy overview similar to enterprise SRS samples."""
        self.doc.add_heading("Visual Overview", level=1)
        items = [
            (self.image_paths.get("system_context"), "Figure A1: System Context"),
            (self.image_paths.get("system_architecture"), "Figure A2: System Architecture"),
            (self.image_paths.get("use_case"), "Figure A3: Use Case Diagram"),
            (self.image_paths.get("user_workflow"), "Figure A4: User Workflow"),
            (self.image_paths.get("security_flow"), "Figure A5: Security Flow"),
            (self.image_paths.get("data_erd"), "Figure A6: Data ERD"),
        ]
        self._add_visual_grid(items, columns=2, image_width=2.9)
    
    def add_external_interfaces_section(
        self, 
        interfaces_data: Dict[str, Any],
        image_paths: Dict[str, str]
    ):
        """
        Add External Interface Requirements section to the document.
        
        Args:
            interfaces_data: Dictionary containing external interfaces section data
            image_paths: Dictionary with paths to interface diagrams
        """
        # Section title
        self.doc.add_heading('9. External Interface Requirements', level=1)
        
        # 9.1 User Interfaces
        user_interfaces = interfaces_data.get('user_interfaces', {})
        self.doc.add_heading(user_interfaces.get('title', '9.1 User Interface'), level=2)
        self.doc.add_paragraph(user_interfaces.get('description', ''))
        
        # Add user interface diagram
        if 'user_interfaces' in image_paths and image_paths['user_interfaces'] and Path(image_paths['user_interfaces']).exists():
            self._add_figure(
                Path(image_paths['user_interfaces']),
                "Figure 7: User Interface Diagram",
                width=6.0,
            )
        
        # 9.2 Hardware / Software / Communication (grouped)
        hardware_interfaces = interfaces_data.get('hardware_interfaces', {})
        self.doc.add_heading(hardware_interfaces.get('title', '9.2 Application Programming Interfaces (APIs)'), level=2)
        self.doc.add_paragraph(hardware_interfaces.get('description', ''))
        
        # Add hardware interface diagram
        if 'hardware_interfaces' in image_paths and image_paths['hardware_interfaces'] and Path(image_paths['hardware_interfaces']).exists():
            self._add_figure(
                Path(image_paths['hardware_interfaces']),
                "Figure 8: Hardware Interface Diagram",
                width=6.0,
            )
        
        # Optional: software/communication diagrams if provided
        if 'software_interfaces' in image_paths and image_paths['software_interfaces'] and Path(image_paths['software_interfaces']).exists():
            self._add_figure(
                Path(image_paths['software_interfaces']),
                "Figure 9: Software/API Interface Diagram",
                width=5.5,
            )
        communication_interfaces = interfaces_data.get('communication_interfaces', {})
        self.doc.add_paragraph(communication_interfaces.get('description', 'REST-based APIs for data communication. Secure API access controls.'))
    
    def add_nfr_section(self, nfr_data: Dict[str, Any]):
        """
        Add Non-Functional Requirements section to the document.
        
        Args:
            nfr_data: Dictionary containing NFR section data
        """
        # Section title
        self.doc.add_heading('6. Non-Functional Requirements', level=1)
        
        # Performance Requirements
        performance = nfr_data.get('performance_requirements', {})
        self.doc.add_heading(f"6.1 {performance.get('title', 'Performance Requirements')}", level=2)
        perf_reqs = performance.get('requirements', [])
        for req in perf_reqs:
            desc = req.get('description', '')
            rationale = req.get('rationale', '')
            para = self.doc.add_paragraph(desc, style='List Bullet')
            if rationale:
                para.add_run(f"\nRationale: {rationale}").italic = True
        
        # Safety Requirements
        safety = nfr_data.get('safety_requirements', {})
        self.doc.add_heading(f"6.2 {safety.get('title', 'Safety Requirements')}", level=2)
        safety_reqs = safety.get('requirements', [])
        for req in safety_reqs:
            desc = req.get('description', '')
            rationale = req.get('rationale', '')
            para = self.doc.add_paragraph(desc, style='List Bullet')
            if rationale:
                para.add_run(f"\nRationale: {rationale}").italic = True
        
        # Security Requirements (brief; full section 7 follows)
        security = nfr_data.get('security_requirements', {})
        self.doc.add_heading(f"6.3 {security.get('title', 'Security Requirements')}", level=2)
        security_reqs = security.get('requirements', [])
        for req in security_reqs:
            desc = req.get('description', '')
            rationale = req.get('rationale', '')
            para = self.doc.add_paragraph(desc, style='List Bullet')
            if rationale:
                para.add_run(f"\nRationale: {rationale}").italic = True
        
        # Quality Attributes
        quality = nfr_data.get('quality_attributes', {})
        self.doc.add_heading(f"6.4 {quality.get('title', 'Quality Attributes')}", level=2)
        quality_reqs = quality.get('requirements', [])
        for req in quality_reqs:
            desc = req.get('description', '')
            rationale = req.get('rationale', '')
            para = self.doc.add_paragraph(desc, style='List Bullet')
            if rationale:
                para.add_run(f"\nRationale: {rationale}").italic = True
    
    def add_security_requirements_section(self):
        """Section 7: Security Requirements with Security Flow diagram."""
        self.doc.add_heading('7. Security Requirements', level=1)
        self.doc.add_paragraph(
            "The system shall: Enforce secure authentication mechanisms; Use encrypted communication channels; "
            "Protect sensitive data from unauthorized access; Log security-related events."
        )
        path = self.image_paths.get('security_flow')
        if path and Path(path).exists():
            self._add_figure(
                Path(path),
                "Figure 5: Security Flow Diagram",
            )
        else:
            self.doc.add_paragraph("Note: The Security Flow Diagram can be included above.", style='Heading 3')
    
    def add_data_requirements_section(self):
        """Section 8: Data Requirements with ERD."""
        self.doc.add_heading('8. Data Requirements', level=1)
        self.doc.add_paragraph("The system shall manage structured data entities and their relationships efficiently.")
        path = self.image_paths.get('data_erd')
        if path and Path(path).exists():
            self._add_figure(
                Path(path),
                "Figure 6: Entity Relationship Diagram",
            )
        else:
            self.doc.add_paragraph("Note: The Entity Relationship Diagram can be included above.", style='Heading 3')
    
    def add_glossary_section(self, glossary_data: Dict[str, Any]):
        """Optional: Glossary (definitions are in 1.3). Kept for backward compatibility."""
        sections = glossary_data.get('sections', [])
        if not sections:
            return
        self.doc.add_heading("Definitions (additional)", level=2)
        for section in sections:
            terms = section.get('terms', [])
            for term_data in terms:
                term = term_data.get('term', '')
                definition = term_data.get('definition', '')
                self.doc.add_paragraph(f"{term} – {definition}", style='List Bullet')
    
    def add_assumptions_section(self, assumptions_data: Dict[str, Any]):
        """
        Add Assumptions and Dependencies (Section 10).
        """
        # Section title
        self.doc.add_heading('10. Assumptions and Dependencies', level=1)
        
        assumptions = assumptions_data.get('assumptions', [])
        for idx, assumption in enumerate(assumptions, 1):
            description = assumption.get('description', '')
            impact = assumption.get('impact', '')
            
            self.doc.add_paragraph(f"Assumption {idx}:", style='Heading 3')
            self.doc.add_paragraph(description)
            
            if impact:
                para = self.doc.add_paragraph()
                para.add_run("Impact: ").bold = True
                para.add_run(impact)
    
    def add_future_enhancements_section(self):
        """Section 11: Future Enhancements."""
        self.doc.add_heading('11. Future Enhancements', level=1)
        for item in [
            "Mobile application support",
            "Advanced analytics and insights",
            "Integration with additional third-party services",
        ]:
            self.doc.add_paragraph(item, style='List Bullet')
    
    def save(self, output_path: str):
        """
        Save the document to the specified path.
        
        Args:
            output_path: Path where the document should be saved
        """
        self.doc.save(output_path)


def generate_srs_document(
    project_name: str,
    introduction_section: Dict[str, Any],
    overall_description_section: Dict[str, Any],
    system_features_section: Dict[str, Any],
    external_interfaces_section: Dict[str, Any],
    nfr_section: Dict[str, Any],
    glossary_section: Dict[str, Any],
    assumptions_section: Dict[str, Any],
    image_paths: Dict[str, str],
    output_path: str,
    authors: List[str] = None,
    organization: str = "Organization Name"
) -> str:
    """
    Generate a complete SRS document from JSON data with Table of Contents.
    
    Args:
        project_name: Name of the project
        introduction_section: Introduction section data
        overall_description_section: Overall description section data
        system_features_section: System features section data
        external_interfaces_section: External interfaces section data
        nfr_section: Non-functional requirements section data
        glossary_section: Glossary section data
        assumptions_section: Assumptions section data
        image_paths: Dictionary with paths to interface diagrams
            Expected keys: 'user_interfaces', 'hardware_interfaces', 
                          'software_interfaces', 'communication_interfaces'
        output_path: Path where the document should be saved
        authors: List of document author names (default: ["Author Name"])
        organization: Organization name (default: "Organization Name")
    
    Returns:
        str: Path to the generated document
    
    Example:
        ```python
        image_paths = {
            'user_interfaces': './static/HireSmart_user_interfaces_diagram.png',
            'hardware_interfaces': './static/HireSmart_hardware_interfaces_diagram.png',
            'software_interfaces': './static/HireSmart_software_interfaces_diagram.png',
            'communication_interfaces': './static/HireSmart_communication_interfaces_diagram.png'
        }
        
        generate_srs_document(
            project_name="HireSmart",
            introduction_section=intro_data,
            overall_description_section=desc_data,
            system_features_section=features_data,
            external_interfaces_section=interfaces_data,
            nfr_section=nfr_data,
            glossary_section=glossary_data,
            assumptions_section=assumptions_data,
            image_paths=image_paths,
            output_path="./output/HireSmart_SRS.docx",
            authors=["John Doe", "Jane Smith"],
            organization="ABC Corporation"
        )
        ```
    """
    # Create generator instance
    generator = SRSDocumentGenerator(project_name, authors, organization)
    generator.image_paths = image_paths
    
    # Add title page
    generator._add_title_page()
    
    # Add Table of Contents
    generator._add_table_of_contents()

    # Add a diagram-heavy visual overview to match enterprise SRS samples
    generator.add_visual_overview_section()

    # Add all sections (proper SRS structure 1–11 with diagram placeholders)
    generator.add_introduction_section(introduction_section)
    generator.add_feasibility_section() # Added Feasibility
    generator.add_overall_description_section(overall_description_section)
    generator.add_system_architecture_section()
    generator.add_system_features_section(system_features_section)
    generator.add_user_workflow_section()
    generator.add_nfr_section(nfr_section)
    generator.add_security_requirements_section()
    generator.add_data_requirements_section()
    generator.add_external_interfaces_section(external_interfaces_section, image_paths)
    generator.add_assumptions_section(assumptions_section)
    generator.add_future_enhancements_section()
    
    # Add header and footer (must be after all content is added)
    generator._apply_layout_to_all_sections()
    generator._add_header_footer()
    generator._set_update_fields_on_open()
    
    # Save the document
    generator.save(output_path)
    
    return output_path
