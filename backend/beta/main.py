from fastapi import FastAPI, Request, HTTPException, BackgroundTasks, Query, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from starlette.concurrency import run_in_threadpool
from fastapi.responses import HTMLResponse, JSONResponse, FileResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from google.adk.sessions import InMemorySessionService
import uuid
from dotenv import load_dotenv, find_dotenv
from pathlib import Path
from backend.beta.agents.introduction_agent import create_introduction_agent 
from backend.beta.agents.overall_description_agent import create_overall_description_agent 
from backend.beta.agents.system_features_agent import create_system_features_agent
from backend.beta.agents.external_interfaces_agent import create_external_interfaces_agent
from backend.beta.agents.nfr_agent import create_nfr_agent
from backend.beta.agents.glossary_agent import create_glossary_agent
from backend.beta.agents.assumptions_agent import create_assumptions_agent
from backend.beta.schemas.srs_input_schema import SRSRequest
from pydantic import BaseModel
from typing import List, Optional
import os
import requests
import base64
import tempfile
import shutil
from docx import Document
from docx.shared import Inches
from backend.beta.utils.email_service import _resolve_docx_path

_root_env = find_dotenv()
_backend_env = Path(__file__).resolve().parents[1] / ".env"
if _backend_env.exists():
    load_dotenv(_backend_env, override=False)
if _root_env:
    load_dotenv(_root_env, override=False)

class NotebookRequest(BaseModel):
    content: str

class ImageRequest(BaseModel):
    prompt: str
    width: int = 1024
    height: int = 1024
    steps: int = 40
    cfg_scale: float = 9.0

class DiagramImageRequest(BaseModel):
    prompt: str

class AppendDiagramRequest(BaseModel):
    content: str
    caption: Optional[str] = None

def _sanitize_mermaid_id(value: str, fallback: str) -> str:
    if not value:
        return fallback
    safe = "".join(ch if ch.isalnum() else "_" for ch in value)
    return safe if safe.strip("_") else fallback

def _reactflow_to_mermaid(nodes: list, edges: list) -> str:
    lines = ["flowchart LR"]
    id_map = {}
    used = set()

    for idx, node in enumerate(nodes or [], start=1):
        raw_id = str(node.get("id") or idx)
        node_id = _sanitize_mermaid_id(raw_id, f"N{idx}")
        if node_id in used:
            node_id = f"{node_id}_{idx}"
        used.add(node_id)
        id_map[raw_id] = node_id

        label = str(node.get("label") or f"Node {idx}").replace('"', "'")
        node_type = str(node.get("type") or "").lower()
        if node_type == "db":
            shape = f'{node_id}[(\"{label}\")]'
        elif node_type == "client":
            shape = f'{node_id}([\"{label}\"])'
        else:
            shape = f'{node_id}[\"{label}\"]'
        lines.append(f"    {shape}")

    if not id_map:
        lines.append("    A[\"System\"] --> B[\"Service\"]")
        return "\n".join(lines)

    for edge in edges or []:
        raw_source = str(edge.get("source") or "")
        raw_target = str(edge.get("target") or "")
        source = id_map.get(raw_source)
        target = id_map.get(raw_target)
        if not source or not target:
            continue
        label = str(edge.get("label") or "").strip()
        if label:
            lines.append(f"    {source} -->|{label}| {target}")
        else:
            lines.append(f"    {source} --> {target}")

    return "\n".join(lines)

async def _generate_diagram_image_data(content: str) -> bytes:
    diagram_payload = {"content": content}
    flow = await WorkflowService.execute_workflow("diagram", diagram_payload)
    mermaid_code = _reactflow_to_mermaid(flow.get("nodes", []), flow.get("edges", []))
    with tempfile.TemporaryDirectory() as tmpdir:
        output_png = Path(tmpdir) / "diagram.png"
        render_mermaid_png(mermaid_code, output_png)
        return output_png.read_bytes()

class NotebookChatRequest(BaseModel):
    content: str
    query: str
    history: List[dict] = []
from backend.beta.utils.globals import (
    create_session ,
    create_runner ,
    create_prompt ,
    generated_response ,
    get_session ,
    clean_and_parse_json,
    clean_interface_diagrams,
    render_mermaid_png)
from google.adk.agents import SequentialAgent , ParallelAgent
from pathlib import Path
import os
import time
from datetime import datetime
from backend.beta.utils.srs_document_generator import generate_srs_document
from backend.beta.utils.model import API_KEY_CONFIGURED, GROQ_API_KEY, GROQ_MODEL, GEMINI_API_KEY
from backend.beta.utils.fallback_srs import build_minimal_sections
from backend.beta.utils.srs_diagrams import get_all_srs_diagrams
import json
from concurrent.futures import ThreadPoolExecutor, as_completed
from litellm import completion as litellm_completion
import threading

today = datetime.today().strftime("%m/%d/%Y")

app = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)

SRS_PROGRESS = {}
SRS_PROGRESS_LOCK = threading.Lock()

app.mount(
    "/static",
    StaticFiles(directory="backend/beta/static"),
    name="static"
)


templates = Jinja2Templates(directory="backend/beta/templates")

session_service_stateful = InMemorySessionService()


async def create_srs_agent():
     
    # Using Parallel execution with Gemini Pro (better rate limits)
    first_agent = SequentialAgent(
          name = "first_agent",
          sub_agents = [
               ParallelAgent(
                    name = "first_parallel_agent",
                    sub_agents = [
                         create_introduction_agent(),
                         create_overall_description_agent(),
                         create_system_features_agent(),
                         create_external_interfaces_agent(),
                         create_nfr_agent()
                    ],
                    description = "This agent handles the generation of the Introduction and Overall Description sections of the SRS document."
               )
          ]
     )

    second_agent = SequentialAgent(
          name = "second_agent",
          sub_agents = [
               ParallelAgent(
                   name = "finalization_agent",
                   sub_agents = [
                       create_glossary_agent(),
                       create_assumptions_agent()
                   ],
                   description = "This agent handles the generation of the Glossary and Assumptions sections of the SRS document."
               )
          ]
     )

    
    return first_agent , second_agent











@app.get("/", response_class=HTMLResponse)
async def root(request: Request):
    return templates.TemplateResponse(
        "home.html",
        {"request": request}
    )


@app.get("/download_srs/{filename}")
async def download_srs(filename: str):
    """Serve the generated SRS .docx for download. Filename e.g. ProjectName_SRS.docx."""
    if not filename.endswith(".docx") or ".." in filename or "/" in filename or "\\" in filename:
        raise HTTPException(status_code=400, detail="Invalid filename")
    base = Path("./backend/beta/generated_srs").resolve()
    path = (base / filename).resolve()
    if not path.is_file() or base not in path.parents:
        raise HTTPException(status_code=404, detail="Document not found")
    return FileResponse(path, filename=filename, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


def _ensure_output_dir():
    Path("./backend/beta/generated_srs").mkdir(exist_ok=True, parents=True)


def _set_progress(project_key: str, stage: str, progress: int, message: str, status: str = "processing", **extra):
    payload = {
        "project_key": project_key,
        "stage": stage,
        "progress": max(0, min(int(progress), 100)),
        "status": status,
        "message": message,
        "updated_at": int(time.time()),
    }
    payload.update(extra)
    with SRS_PROGRESS_LOCK:
        SRS_PROGRESS[project_key] = payload


def _get_progress(project_key: str) -> dict:
    with SRS_PROGRESS_LOCK:
        return SRS_PROGRESS.get(project_key, {
            "project_key": project_key,
            "stage": "idle",
            "progress": 0,
            "status": "idle",
            "message": "No generation started for this project.",
            "updated_at": int(time.time()),
        })


def _safe_project_key(project_name: str) -> str:
    safe = "".join(ch if ch.isalnum() or ch in ("-", "_") else "_" for ch in (project_name or "Project"))
    safe = safe.strip("_")
    return safe or "Project"


def _build_image_paths(project_key: str) -> dict:
    return {
        "user_interfaces": Path(f"./backend/beta/static/{project_key}_user_interfaces_diagram.png"),
        "hardware_interfaces": Path(f"./backend/beta/static/{project_key}_hardware_interfaces_diagram.png"),
        "software_interfaces": Path(f"./backend/beta/static/{project_key}_software_interfaces_diagram.png"),
        "communication_interfaces": Path(f"./backend/beta/static/{project_key}_communication_interfaces_diagram.png"),
        "system_context": Path(f"./backend/beta/static/{project_key}_system_context.png"),
        "system_architecture": Path(f"./backend/beta/static/{project_key}_system_architecture.png"),
        "use_case": Path(f"./backend/beta/static/{project_key}_use_case.png"),
        "user_workflow": Path(f"./backend/beta/static/{project_key}_user_workflow.png"),
        "security_flow": Path(f"./backend/beta/static/{project_key}_security_flow.png"),
        "data_erd": Path(f"./backend/beta/static/{project_key}_data_erd.png"),
        "sequence_diagram": Path(f"./backend/beta/static/{project_key}_sequence_diagram.png"),
        "state_diagram": Path(f"./backend/beta/static/{project_key}_state_diagram.png"),
        "ui_local_diagram": Path(f"./backend/beta/static/{project_key}_ui_local_diagram.png"),
    }


def _output_path(project_key: str, variant: str = "full") -> str:
    if variant == "instant":
        return f"./backend/beta/generated_srs/{project_key}_SRS_instant.docx"
    if variant == "quick":
        return f"./backend/beta/generated_srs/{project_key}_SRS_quick.docx"
    if variant == "enhanced":
        return f"./backend/beta/generated_srs/{project_key}_SRS_enhanced.docx"
    return f"./backend/beta/generated_srs/{project_key}_SRS.docx"


def _to_list(value):
    if value is None:
        return []
    if isinstance(value, list):
        return [str(v).strip() for v in value if str(v).strip()]
    return [str(value).strip()] if str(value).strip() else []


def _select_gemini_models() -> list:
    """
    Return ordered model candidates from fastest to richest.
    Environment override:
    - GEMINI_MODEL: comma-separated list, e.g. "gemini-2.0-flash,gemini-1.5-pro-latest"
    """
    raw = os.getenv("GEMINI_MODEL", "").strip()
    if raw:
        return [m.strip() for m in raw.split(",") if m.strip()]
    return [
        "gemini-2.0-flash",
        "gemini-1.5-flash",
        "gemini-1.5-pro-latest",
    ]


def _try_fast_litellm_json(prompt: str) -> dict:
    """
    Try Groq/LiteLLM first for speed. Returns parsed JSON dict or {}.
    """
    try:
        if not GROQ_API_KEY:
            return {}
        # Prefer configured GROQ model first if it is not gemini/*
        model_candidates = []
        if GROQ_MODEL and not GROQ_MODEL.startswith("gemini/"):
            model_candidates.append(GROQ_MODEL)
        # Fast Groq defaults
        model_candidates.extend([
            "groq/llama-3.1-8b-instant",
            "groq/llama-3.3-70b-versatile",
        ])
        # Expand Groq model aliases so both raw and provider-qualified names are tried.
        expanded = []
        for m in model_candidates:
            expanded.append(m)
            if not m.startswith("groq/"):
                expanded.append(f"groq/{m}")

        # Deduplicate while keeping order
        seen = set()
        models = [m for m in expanded if not (m in seen or seen.add(m))]

        for model_name in models:
            try:
                print(f"⚡ Trying fast LiteLLM model: {model_name}")
                # Some models/providers do not support strict response_format json_object.
                # We retry without it if needed.
                try:
                    resp = litellm_completion(
                        model=model_name,
                        messages=[{"role": "user", "content": prompt}],
                        temperature=0,
                        response_format={"type": "json_object"},
                        timeout=15,
                    )
                except Exception:
                    resp = litellm_completion(
                        model=model_name,
                        messages=[{"role": "user", "content": prompt}],
                        temperature=0,
                        timeout=15,
                    )
                text = (resp.choices[0].message.content or "").strip()
                cleaned = clean_and_parse_json(text)
                if cleaned:
                    print(f"✅ LiteLLM JSON accepted from: {model_name}")
                    return cleaned
                print(f"⚠️ LiteLLM returned non-parseable JSON from: {model_name}")
            except Exception as e:
                print(f"⚠️ LiteLLM model failed ({model_name}): {e}")
        return {}
    except Exception as e:
        print(f"⚠️ LiteLLM fast path failed: {e}")
        return {}


def _map_ai_to_sections(inputs: dict, ai: dict) -> dict:
    """Convert flexible AI output into strict generator section structure."""
    base = build_minimal_sections(inputs)
    if not isinstance(ai, dict) or not ai:
        return base

    sections = {
        "introduction_section": dict(base["introduction_section"]),
        "overall_description_section": dict(base["overall_description_section"]),
        "system_features_section": dict(base["system_features_section"]),
        "external_interfaces_section": dict(base["external_interfaces_section"]),
        "nfr_section": dict(base["nfr_section"]),
        "glossary_section": dict(base["glossary_section"]),
        "assumptions_section": dict(base["assumptions_section"]),
    }

    intro_ai = ai.get("introduction", {}) if isinstance(ai.get("introduction"), dict) else {}
    purpose_ai = intro_ai.get("purpose")
    if isinstance(purpose_ai, str) and purpose_ai.strip():
        sections["introduction_section"]["purpose"]["description"] = purpose_ai.strip()
    elif isinstance(purpose_ai, dict):
        desc = purpose_ai.get("description")
        if isinstance(desc, str) and desc.strip():
            sections["introduction_section"]["purpose"]["description"] = desc.strip()

    scope_ai = intro_ai.get("scope") if isinstance(intro_ai.get("scope"), dict) else {}
    if scope_ai:
        desc = scope_ai.get("description")
        if isinstance(desc, str) and desc.strip():
            sections["introduction_section"]["project_scope"]["description"] = desc.strip()
        included = _to_list(scope_ai.get("included"))
        excluded = _to_list(scope_ai.get("excluded"))
        if included:
            sections["introduction_section"]["project_scope"]["included"] = included
        if excluded:
            sections["introduction_section"]["project_scope"]["excluded"] = excluded

    definitions = intro_ai.get("definitions")
    if isinstance(definitions, list):
        terms = []
        for item in definitions:
            if isinstance(item, dict):
                term = str(item.get("term", "")).strip()
                definition = str(item.get("definition", "")).strip()
                if term and definition:
                    terms.append({"term": term, "definition": definition})
        if terms:
            sections["glossary_section"] = {"sections": [{"title": "Definitions", "terms": terms}]}

    overall_ai = ai.get("overall_description", {}) if isinstance(ai.get("overall_description"), dict) else {}
    perspective = overall_ai.get("product_perspective")
    if isinstance(perspective, str) and perspective.strip():
        sections["overall_description_section"]["product_perspective"]["description"] = perspective.strip()

    user_chars = overall_ai.get("user_characteristics")
    if isinstance(user_chars, list):
        classes = []
        for item in user_chars:
            if isinstance(item, dict):
                user_class = str(item.get("user_class", "")).strip()
                if user_class:
                    classes.append({
                        "user_class": user_class,
                        "characteristics": str(item.get("characteristics", "")).strip(),
                        "responsibilities": str(item.get("responsibilities", "")).strip(),
                        "skills": str(item.get("skills", "")).strip()
                    })
        if classes:
            sections["overall_description_section"]["user_classes_and_characteristics"]["user_classes"] = classes

    assumptions = _to_list(overall_ai.get("assumptions"))
    if assumptions:
        sections["assumptions_section"]["assumptions"] = [{"description": a, "impact": ""} for a in assumptions]

    fr_ai = ai.get("functional_requirements")
    if isinstance(fr_ai, list):
        features = []
        for item in fr_ai:
            if not isinstance(item, dict):
                continue
            feature_name = str(item.get("feature_name", "")).strip()
            if not feature_name:
                continue
            description = str(item.get("description", "")).strip()
            requirements = _to_list(item.get("requirements"))
            structured = item.get("structured_requirements", {})
            
            if not requirements:
                requirements = [f"Support: {feature_name}"]
            
            features.append({
                "feature_name": feature_name,
                "description": description,
                "functional_requirements": [{"description": req} for req in requirements],
                "structured_requirements": structured if isinstance(structured, dict) else {}
            })
        if features:
            sections["system_features_section"]["features"] = features

    nfr_ai = ai.get("non_functional_requirements", {}) if isinstance(ai.get("non_functional_requirements"), dict) else {}
    perf = _to_list(nfr_ai.get("performance"))
    security = _to_list(nfr_ai.get("security"))
    reliability = _to_list(nfr_ai.get("reliability"))
    if perf:
        sections["nfr_section"]["performance_requirements"]["requirements"] = [{"description": p} for p in perf]
    if security:
        sections["nfr_section"]["security_requirements"]["requirements"] = [{"description": s} for s in security]
    if reliability:
        sections["nfr_section"]["safety_requirements"]["requirements"] = [{"description": r} for r in reliability]

    risk_ai = ai.get("risk_analysis")
    if isinstance(risk_ai, list):
        sections["risk_analysis"] = [r for r in risk_ai if isinstance(r, dict)]

    return sections


def _render_all_diagrams(inputs: dict, image_paths: dict, interface_sections: dict):
    """Render all major and interface Mermaid diagrams to PNG files."""
    stats = {"core_rendered": 0, "core_failed": 0, "interface_rendered": 0, "interface_failed": 0}
    interface_keys = [
        "user_interfaces",
        "hardware_interfaces",
        "software_interfaces",
        "communication_interfaces",
    ]
    render_jobs = []

    # Core diagrams
    for key, mermaid_code in get_all_srs_diagrams(inputs).items():
        output_png = image_paths.get(key)
        if output_png and isinstance(mermaid_code, str) and mermaid_code.strip():
            render_jobs.append((key, mermaid_code, output_png, "core"))

    # Interface diagrams
    for key in interface_keys:
        section = interface_sections.get(key, {})
        code = ((section.get("interface_diagram") or {}).get("code", "") if isinstance(section, dict) else "")
        output_png = image_paths.get(key)
        if output_png and isinstance(code, str) and code.strip():
            render_jobs.append((key, code, output_png, "interface"))

    # Parallel rendering for faster response time.
    max_workers = max(2, min(6, len(render_jobs)))
    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        future_map = {
            executor.submit(render_mermaid_png, code, output_png): (key, kind)
            for key, code, output_png, kind in render_jobs
        }
        for future in as_completed(future_map):
            key, kind = future_map[future]
            try:
                future.result()
                if kind == "core":
                    stats["core_rendered"] += 1
                else:
                    stats["interface_rendered"] += 1
            except Exception as e:
                if kind == "core":
                    stats["core_failed"] += 1
                    print(f"⚠️ Failed to render {key} diagram: {e}")
                else:
                    stats["interface_failed"] += 1
                    print(f"⚠️ Failed to render {key} interface diagram: {e}")
    return stats


def _render_quick_diagrams(inputs: dict, image_paths: dict):
    """Render only 2 core diagrams for quick mode."""
    stats = {"core_rendered": 0, "core_failed": 0, "interface_rendered": 0, "interface_failed": 0}
    all_diagrams = get_all_srs_diagrams(inputs)
    for key in ["system_context", "system_architecture"]:
        code = all_diagrams.get(key)
        output_png = image_paths.get(key)
        if not code or not output_png:
            continue
        try:
            render_mermaid_png(code, output_png)
            stats["core_rendered"] += 1
        except Exception as e:
            stats["core_failed"] += 1
            print(f"⚠️ Failed quick render for {key}: {e}")
    return stats


def _build_sections_with_ai(inputs: dict, project_name: str, project_key: str = "", mode: str = "full") -> dict:
    """Build merged sections using AI when available; fallback to minimal."""
    ai_content = {}
    budget_sec = float(os.getenv("QUICK_AI_BUDGET_SEC", "18")) if mode == "quick" else float(os.getenv("FULL_AI_BUDGET_SEC", "90"))
    started = time.monotonic()

    if API_KEY_CONFIGURED:
        try:
            print(f"🚀 Starting AI Expansion for: {project_name}")
            if project_key:
                _set_progress(project_key, "ai", 25, "Generating detailed requirements with AI...")
            import google.generativeai as genai
            
            genai.configure(api_key=os.getenv("GEMINI_API_KEY"))
            
            detail_instruction = "Ensure the content is concise but professional."
            extra_instructions = (inputs.get("output_control") or {}).get("additional_instructions")
            if extra_instructions:
                detail_instruction = f"{detail_instruction}\nAdditional instructions: {extra_instructions}"
            json_structure = """
                "functional_requirements": [
                    {
                        "feature_name": "Feature Name",
                        "description": "2-3 feature description.",
                        "requirements": ["Req 1", "Req 2", "Req 3"]
                    }
                ],
                "overall_description": {
                    "product_perspective": "...",
                    "user_characteristics": [
                        {"user_class": "Admin", "characteristics": "..."} 
                    ],
                    "assumptions": [...]
                }
            """

            if mode in ["full", "enhanced"]:
                detail_instruction = (
                    "PROVIDE EXTENSIVE ENTERPRISE-GRADE DETAIL. "
                    "Write comprehensive, professional paragraphs (100-150 words) for descriptions. "
                    "Use structured data for tables."
                )
                json_structure = """
                "functional_requirements": [
                    {
                        "feature_name": "Feature Name",
                        "description": "Detailed description of the feature.",
                        "requirements": ["Req 1", "Req 2"],
                        "structured_requirements": {
                             "inputs": "User ID, Password...",
                             "outputs": "Dashboard, Error Message...",
                             "acceptance_criteria": "User must be redirected within 2s..."
                        }
                    }
                ],
                "overall_description": {
                    "product_perspective": "Detailed 200-word perspective...",
                    "user_characteristics": [
                        {
                            "user_class": "Admin", 
                            "characteristics": "System administrator...",
                            "responsibilities": "System config, User management",
                            "skills": "High technical proficiency"
                        } 
                    ],
                    "assumptions": [...]
                },
                "risk_analysis": [
                    {
                        "risk": "Data Breach",
                        "probability": "Low",
                        "impact": "High",
                        "mitigation": "Encryption at rest..."
                    }
                ]
                """

            prompt = f"""
            You are an expert Senior Technical Writer. I need you to generate a comprehensive IEEE 830 Software Requirements Specification (SRS) in JSON format.
            
            {detail_instruction}
            
            Project Input Data:
            {json.dumps(inputs, indent=2)}
            
            Task:
            Expand the short user inputs into detailed, professional technical content.
            
            Required JSON Structure:
            {{
                "introduction": {{
                    "purpose": "50-75 word professional summary of the system purpose.",
                    "scope": {{
                        "description": "100 word description of what the system does.",
                        "included": ["List of 5-7 in-scope features/modules"],
                        "excluded": ["List of 3 out-of-scope items"]
                    }},
                    "definitions": [
                         {{"term": "Term1", "definition": "Def1"}}
                    ]
                }},
                {json_structure},
                "non_functional_requirements": {{
                    "performance": ["Req 1", "Req 2"],
                    "security": ["Req 1", "Req 2"],
                    "reliability": ["Req 1"]
                }}
            }}
            
            Constraints:
            - Generate at least 5 functional requirements when possible.
            - Return ONLY valid JSON (no markdown, no comments).
            """

            # Fast path first (Groq/LiteLLM), then Gemini fallback
            ai_content = _try_fast_litellm_json(prompt)
            if ai_content:
                if project_key:
                    _set_progress(project_key, "ai", 40, "AI content generated (fast provider).")
                sections = _map_ai_to_sections(inputs, ai_content)
                interface_sections = clean_interface_diagrams(sections.get("external_interfaces_section", {}))
                sections["external_interfaces_section"] = interface_sections
                return sections

            for model_name in _select_gemini_models():
                if (time.monotonic() - started) > budget_sec:
                    print(f"⏱️ AI budget exceeded ({budget_sec}s). Using fallback content.")
                    break
                try:
                    print(f"⚡ Trying model: {model_name}")
                    model = genai.GenerativeModel(model_name)
                    response = model.generate_content(prompt)
                    cleaned_json = clean_and_parse_json(response.text)
                    if cleaned_json:
                        ai_content = cleaned_json
                        print(f"✅ AI Response accepted from: {model_name}")
                        if project_key:
                            _set_progress(project_key, "ai", 40, f"AI content generated ({model_name}).")
                        break
                    print(f"⚠️ Parsed empty JSON from model: {model_name}")
                except Exception as model_err:
                    print(f"⚠️ Model {model_name} failed: {model_err}")
            if not ai_content:
                print("⚠️ All models failed or returned invalid JSON. Falling back.")
                
        except Exception as e:
            print(f"⚠️ AI Expansion failed: {e}")
            import traceback
            traceback.print_exc()

    sections = _map_ai_to_sections(inputs, ai_content)
    if not ai_content:
        print("ℹ️ Using Minimal Fallback")
        if project_key:
            _set_progress(project_key, "ai", 35, "Using fallback baseline content.")
    interface_sections = clean_interface_diagrams(sections.get("external_interfaces_section", {}))
    sections["external_interfaces_section"] = interface_sections
    return sections


def _generate_document(project_name: str, project_key: str, inputs: dict, sections: dict, image_paths: dict, variant: str):
    output_file = generate_srs_document(
        project_name=project_name,
        introduction_section=sections["introduction_section"],
        overall_description_section=sections["overall_description_section"],
        system_features_section=sections["system_features_section"],
        external_interfaces_section=sections["external_interfaces_section"],
        nfr_section=sections["nfr_section"],
        glossary_section=sections["glossary_section"],
        assumptions_section=sections["assumptions_section"],
        image_paths=image_paths,
        output_path=_output_path(project_key, variant),
        authors=inputs["project_identity"]["author"],
        organization=inputs["project_identity"]["organization"],
        sections=sections,
        mode=variant
    )
    
    # Auto-publish as sample report for Landing Page demo
    try:
        import shutil
        sample_path = os.path.join("backend", "beta", "static", "sample_report.docx")
        shutil.copy(output_file, sample_path)
        print(f"📄 Updated public sample report at: {sample_path}")
    except Exception as e:
        print(f"⚠️ Failed to update sample report: {e}")
        
    return output_file


def _generate_instant_fallback(project_name: str, project_key: str, inputs: dict, image_paths: dict):
    """
    Guaranteed low-latency fallback when quick/full document composition fails.
    Uses minimal sections and only pre-existing core images.
    """
    sections = build_minimal_sections(inputs)
    sections["external_interfaces_section"] = clean_interface_diagrams(sections.get("external_interfaces_section", {}))
    instant_image_paths = {}
    for k in ["system_context", "system_architecture"]:
        p = image_paths.get(k)
        if p and Path(p).is_file():
            instant_image_paths[k] = p
    return _generate_document(project_name, project_key, inputs, sections, instant_image_paths, "instant")


def _generate_enhanced_background(inputs: dict, project_name: str, project_key: str):
    """Background task to create enhanced SRS after quick file is returned."""
    try:
        print(f"🛠️ Background enhanced generation started: {project_name}")
        _set_progress(project_key, "enhanced_ai", 88, "Preparing enhanced version...", status="processing")
        image_paths = _build_image_paths(project_key)
        sections = _build_sections_with_ai(inputs, project_name, project_key)
        _set_progress(project_key, "enhanced_diagrams", 92, "Rendering enhanced diagrams...", status="processing")
        diagram_stats = _render_all_diagrams(inputs, image_paths, sections["external_interfaces_section"])
        if diagram_stats["core_rendered"] == 0:
            _set_progress(project_key, "enhanced_diagrams", 93, "Diagrams unavailable; continuing enhanced build.", status="processing")
        _set_progress(project_key, "enhanced_doc", 96, "Compiling enhanced DOCX...", status="processing")
        _generate_document(project_name, project_key, inputs, sections, image_paths, "enhanced")
        _set_progress(project_key, "completed", 100, "Enhanced document ready.", status="completed")
        print(f"✅ Background enhanced generation completed: {project_name}")
    except Exception as e:
        _set_progress(project_key, "failed", 100, f"Enhanced generation failed: {e}", status="failed")
        print(f"❌ Background enhanced generation failed for {project_name}: {e}")


@app.get("/srs_status/{project_key}")
async def srs_status(project_key: str):
    """Check quick/enhanced document readiness for two-phase generation."""
    instant_path = Path(_output_path(project_key, "instant"))
    quick_path = Path(_output_path(project_key, "quick"))
    enhanced_path = Path(_output_path(project_key, "enhanced"))
    full_path = Path(_output_path(project_key, "full"))
    return {
        "project_key": project_key,
        "instant_ready": instant_path.is_file(),
        "quick_ready": quick_path.is_file(),
        "enhanced_ready": enhanced_path.is_file(),
        "full_ready": full_path.is_file(),
        "instant_download_url": f"/download_srs/{instant_path.name}" if instant_path.is_file() else None,
        "quick_download_url": f"/download_srs/{quick_path.name}" if quick_path.is_file() else None,
        "enhanced_download_url": f"/download_srs/{enhanced_path.name}" if enhanced_path.is_file() else None,
        "full_download_url": f"/download_srs/{full_path.name}" if full_path.is_file() else None,
    }


@app.get("/srs_progress/{project_key}")
async def srs_progress(project_key: str):
    """Stage-wise progress for SRS generation."""
    return _get_progress(project_key)


@app.post("/generate_srs")
async def generate_srs(
    srs_data: SRSRequest,
    background_tasks: BackgroundTasks,
    mode: str = Query(default="full", regex="^(full|quick|instant)$"),
):
    inputs = srs_data.dict()
    project_name = inputs["project_identity"]["project_name"]
    # If project_id is provided from Node.js, use it as the stable key for files & progress.
    # This prevents creating multiple files for the same project during regeneration.
    project_id = inputs["project_identity"].get("project_id")
    if project_id:
        project_key = str(project_id)
    else:
        project_key = _safe_project_key(project_name)
    
    _ensure_output_dir()
    image_paths = _build_image_paths(project_key)
    _set_progress(project_key, "init", 5, "Initializing generation...")

    try:
        if mode == "instant":
            # Fastest path: no live Mermaid rendering in request cycle.
            _set_progress(project_key, "content", 20, "Preparing baseline content...")
            sections = build_minimal_sections(inputs)
            sections["external_interfaces_section"] = clean_interface_diagrams(sections.get("external_interfaces_section", {}))
            instant_image_paths = {}
            for k in ["system_context", "system_architecture"]:
                p = image_paths.get(k)
                if p and Path(p).is_file():
                    instant_image_paths[k] = p

            try:
                _set_progress(project_key, "doc", 70, "Building DOCX file...")
                generated_path = await run_in_threadpool(
                    _generate_document,
                    project_name, project_key, inputs, sections, instant_image_paths, "instant"
                )
                _set_progress(
                    project_key,
                    "completed",
                    100,
                    "Instant document ready.",
                    status="completed",
                    download_url=f"/download_srs/{Path(generated_path).name}",
                    mode="instant",
                )
                background_tasks.add_task(_generate_enhanced_background, inputs, project_name, project_key)
                return {
                    "status": "success",
                    "mode": "instant",
                    "message": "Instant SRS generated. Enhanced version is being prepared in background.",
                    "srs_document_path": generated_path,
                    "download_url": f"/download_srs/{Path(generated_path).name}",
                    "enhanced_status_url": f"/srs_status/{project_key}",
                    "enhanced_download_url": f"/download_srs/{Path(_output_path(project_key, 'enhanced')).name}",
                }
            except Exception as e:
                print(f"❌ Instant Document Generation Failed: {e}")
                import traceback
                traceback.print_exc()
                _set_progress(project_key, "failed", 100, str(e), status="failed")
                raise HTTPException(status_code=500, detail=str(e))

        if mode == "quick":
            # Quick mode: AI-enriched sections + only 2 core diagrams (better quality, faster than full).
            sections = _build_sections_with_ai(inputs, project_name, project_key)
            _set_progress(project_key, "diagrams", 55, "Rendering core diagrams...")
            quick_stats = await run_in_threadpool(_render_quick_diagrams, inputs, image_paths)
            if quick_stats["core_rendered"] == 0:
                # Do not fail quick mode; generate document without freshly rendered diagrams.
                _set_progress(
                    project_key,
                    "diagrams",
                    60,
                    "Core diagrams unavailable; continuing with document generation.",
                    status="processing",
                )
            try:
                _set_progress(project_key, "doc", 80, "Compiling quick DOCX...")
                generated_path = await run_in_threadpool(
                    _generate_document, project_name, project_key, inputs, sections, image_paths, "quick"
                )
                _set_progress(
                    project_key,
                    "completed",
                    100,
                    "Quick document ready.",
                    status="completed",
                    download_url=f"/download_srs/{Path(generated_path).name}",
                    mode="quick",
                )
                background_tasks.add_task(_generate_enhanced_background, inputs, project_name, project_key)
                return {
                    "status": "success",
                    "mode": "quick",
                    "message": "Quick SRS generated. Enhanced version is being prepared in background.",
                    "srs_document_path": generated_path,
                    "download_url": f"/download_srs/{Path(generated_path).name}",
                    "enhanced_status_url": f"/srs_status/{project_key}",
                    "enhanced_download_url": f"/download_srs/{Path(_output_path(project_key, 'enhanced')).name}",
                    "warnings": [] if quick_stats["core_rendered"] > 0 else [
                        "Core diagrams could not be freshly rendered in quick mode; document was generated with available assets."
                    ],
                }
            except Exception as e:
                print(f"❌ Quick Document Generation Failed: {e}")
                import traceback
                traceback.print_exc()
                try:
                    _set_progress(project_key, "doc", 82, "Quick build failed, switching to instant fallback...")
                    generated_path = await run_in_threadpool(_generate_instant_fallback, project_name, project_key, inputs, image_paths)
                    _set_progress(
                        project_key,
                        "completed",
                        100,
                        "Instant fallback document ready.",
                        status="completed",
                        download_url=f"/download_srs/{Path(generated_path).name}",
                        mode="instant",
                    )
                    background_tasks.add_task(_generate_enhanced_background, inputs, project_name, project_key)
                    return {
                        "status": "success",
                        "mode": "instant",
                        "message": "Quick generation failed, instant fallback generated successfully.",
                        "srs_document_path": generated_path,
                        "download_url": f"/download_srs/{Path(generated_path).name}",
                        "enhanced_status_url": f"/srs_status/{project_key}",
                        "enhanced_download_url": f"/download_srs/{Path(_output_path(project_key, 'enhanced')).name}",
                        "warnings": [f"Quick generation failed: {e}. Returned instant fallback."],
                    }
                except Exception as fallback_err:
                    _set_progress(project_key, "failed", 100, str(fallback_err), status="failed")
                    raise HTTPException(status_code=500, detail=f"Quick and instant fallback failed: {fallback_err}")

        sections = _build_sections_with_ai(inputs, project_name, project_key)
        _set_progress(project_key, "diagrams", 60, "Rendering all diagrams...")
        diagram_stats = await run_in_threadpool(_render_all_diagrams, inputs, image_paths, sections["external_interfaces_section"])
        if diagram_stats["core_rendered"] == 0:
            _set_progress(
                project_key,
                "diagrams",
                70,
                "Diagrams unavailable; continuing with document generation.",
                status="processing",
            )

        # 3. Document Construction Phase
        try:
            _set_progress(project_key, "doc", 85, "Building full DOCX...")
            generated_path = await run_in_threadpool(
                _generate_document, project_name, project_key, inputs, sections, image_paths, "full"
            )
            _set_progress(
                project_key,
                "completed",
                100,
                "Full document ready.",
                status="completed",
                download_url=f"/download_srs/{Path(generated_path).name}",
                mode="full",
            )
            
            return {
                "status": "success",
                "message": "SRS document generated successfully",
                "mode": "full",
                "srs_document_path": generated_path,
                "download_url": f"/download_srs/{Path(generated_path).name}",
                "warnings": [] if diagram_stats["core_rendered"] > 0 else [
                    "Some diagrams could not be rendered; document was generated with available assets."
                ],
            }

        except Exception as e:
            print(f"❌ Document Generation Failed: {e}")
            import traceback
            traceback.print_exc()
            try:
                _set_progress(project_key, "doc", 88, "Full build failed, switching to instant fallback...")
                generated_path = await run_in_threadpool(_generate_instant_fallback, project_name, project_key, inputs, image_paths)
                _set_progress(
                    project_key,
                    "completed",
                    100,
                    "Instant fallback document ready.",
                    status="completed",
                    download_url=f"/download_srs/{Path(generated_path).name}",
                    mode="instant",
                )
                background_tasks.add_task(_generate_enhanced_background, inputs, project_name, project_key)
                return {
                    "status": "success",
                    "message": "Full generation failed, instant fallback generated successfully.",
                    "mode": "instant",
                    "srs_document_path": generated_path,
                    "download_url": f"/download_srs/{Path(generated_path).name}",
                    "warnings": [f"Full generation failed: {e}. Returned instant fallback."],
                }
            except Exception as fallback_err:
                _set_progress(project_key, "failed", 100, str(fallback_err), status="failed")
                raise HTTPException(status_code=500, detail=f"Full and instant fallback failed: {fallback_err}")
    except HTTPException:
        raise
    except Exception as e:
        # Last-resort safety: attempt instant fallback for any unexpected failure.
        try:
            _set_progress(project_key, "doc", 75, "Unexpected error, switching to instant fallback...")
            generated_path = _generate_instant_fallback(project_name, project_key, inputs, image_paths)
            _set_progress(
                project_key,
                "completed",
                100,
                "Instant fallback document ready.",
                status="completed",
                download_url=f"/download_srs/{Path(generated_path).name}",
                mode="instant",
            )
            background_tasks.add_task(_generate_enhanced_background, inputs, project_name, project_key)
            return {
                "status": "success",
                "message": "Unexpected error; instant fallback generated successfully.",
                "mode": "instant",
                "srs_document_path": generated_path,
                "download_url": f"/download_srs/{Path(generated_path).name}",
                "warnings": [f"Unexpected error: {e}. Returned instant fallback."],
            }
        except Exception as fallback_err:
            _set_progress(project_key, "failed", 100, str(fallback_err), status="failed")
            raise HTTPException(status_code=500, detail=f"Unexpected error; instant fallback failed: {fallback_err}")


# --- AI Notebook Endpoints ---

from backend.beta.services.workflow_service import WorkflowService

@app.post("/api/notebook/analyze")
async def analyze_notebook(request: NotebookRequest):
    """
    Analyzes notebook content via WorkflowService (n8n or Local Gemini).
    """
    try:
        payload = {"content": request.content}
        result = await WorkflowService.execute_workflow("analyze", payload)
        return JSONResponse(content=result)
    except Exception as e:
        print(f"Notebook Analysis Error: {e}")
        return JSONResponse(content={"services": []})

@app.post("/api/notebook/chat")
async def chat_notebook(request: NotebookChatRequest):
    """
    Context-aware chat via WorkflowService (n8n or Local Gemini).
    """
    try:
        payload = {
            "content": request.content,
            "query": request.query,
            "history": request.history
        }
        result = await WorkflowService.execute_workflow("chat", payload)
        return JSONResponse(content=result)

    except Exception as e:
        print(f"Notebook Chat Error: {e}")
        return JSONResponse(content={"answer": f"Error interacting with workflow engine: {str(e)}"})

@app.post("/api/notebook/diagram")
async def generate_diagram(request: NotebookRequest):
    """
    Generates ReactFlow diagram JSON via WorkflowService.
    """
    try:
        payload = {"content": request.content}
        result = await WorkflowService.execute_workflow("diagram", payload)
        return JSONResponse(content=result)
    except Exception as e:
        print(f"Diagram Gen Error: {e}")
        return JSONResponse(content={"nodes": [], "edges": []})

@app.post("/api/notebook/image")
async def generate_image(request: ImageRequest):
    api_key = os.getenv("STABILITY_API_KEY", "").strip()
    if not api_key:
        raise HTTPException(status_code=500, detail="STABILITY_API_KEY not configured")

    negative_prompt = os.getenv(
        "STABILITY_NEGATIVE_PROMPT",
        "blurry, distorted, unreadable text, misspelled text, watermark, logo, noisy background"
    ).strip()
    payload = {
        "text_prompts": [
            {"text": request.prompt, "weight": 1},
            {"text": negative_prompt, "weight": -1}
        ],
        "cfg_scale": request.cfg_scale,
        "height": request.height,
        "width": request.width,
        "samples": 1,
        "steps": request.steps
    }
    style_preset = os.getenv("STABILITY_STYLE_PRESET", "line-art").strip()
    if style_preset:
        payload["style_preset"] = style_preset
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json"
    }
    try:
        custom_url = os.getenv("STABILITY_API_URL", "").strip()
        base_url = os.getenv("STABILITY_API_BASE", "https://api.stability.ai/v1/generation").strip()
        engine = os.getenv("STABILITY_ENGINE", "stable-diffusion-xl-1024-v1-0").strip()

        engine_candidates = [engine, "stable-diffusion-xl-1024-v1-0", "stable-diffusion-v1-6"]
        tried = set()
        last_error = None

        for engine_id in engine_candidates:
            if engine_id in tried:
                continue
            tried.add(engine_id)
            api_url = custom_url or f"{base_url}/{engine_id}/text-to-image"
            resp = requests.post(api_url, headers=headers, json=payload, timeout=60)
            if resp.status_code != 200:
                last_error = resp.text
                if resp.status_code in (400, 404) and "not found" in resp.text.lower() and not custom_url:
                    continue
                raise HTTPException(status_code=500, detail=f"Image generation failed: {resp.text}")
            data = resp.json()
            base64_img = None
            if isinstance(data, dict) and data.get("artifacts"):
                base64_img = data["artifacts"][0].get("base64")
            if not base64_img:
                return JSONResponse(content={"image": None})
            return JSONResponse(content={"image": f"data:image/png;base64,{base64_img}"})

        raise HTTPException(status_code=500, detail=f"Image generation failed: {last_error}")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Image generation error: {e}")

@app.get("/api/notebook/diagram-image/status")
async def diagram_image_status():
    mmdc_path = shutil.which("mmdc") or shutil.which("mmdc.cmd")
    return JSONResponse(content={
        "enabled": bool(mmdc_path),
        "reason": "" if mmdc_path else "mmdc not found. Install @mermaid-js/mermaid-cli."
    })

@app.get("/api/notebook/image/status")
async def image_status():
    api_key = os.getenv("STABILITY_API_KEY", "").strip()
    return JSONResponse(content={
        "enabled": bool(api_key),
        "reason": "" if api_key else "STABILITY_API_KEY missing"
    })

@app.post("/api/notebook/diagram-image")
async def generate_diagram_image(request: DiagramImageRequest):
    if not request.prompt.strip():
        raise HTTPException(status_code=400, detail="Prompt is required")
    try:
        data = await _generate_diagram_image_data(request.prompt)
        return JSONResponse(content={"image": f"data:image/png;base64,{base64.b64encode(data).decode('utf-8')}"})
    except FileNotFoundError as e:
        raise HTTPException(status_code=500, detail=str(e))
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Diagram image generation failed: {e}")

@app.post("/api/project/{project_id}/append-diagram")
async def append_diagram_to_project(project_id: str, request: AppendDiagramRequest):
    project = ProjectStore.get_project(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    if not request.content.strip():
        raise HTTPException(status_code=400, detail="Content is required")

    try:
        data = await _generate_diagram_image_data(request.content)
        diagrams_dir = Path("backend/beta/static/diagrams")
        diagrams_dir.mkdir(parents=True, exist_ok=True)
        filename = f"{project_id}_diagram_{int(time.time())}.png"
        image_path = diagrams_dir / filename
        image_path.write_bytes(data)

        image_url = f"/static/diagrams/{filename}"
        caption = (request.caption or "Studio Diagram").strip()
        markdown_block = f"\n\n## {caption}\n\n![{caption}]({image_url})\n"
        project.contentMarkdown = (project.contentMarkdown or "") + markdown_block

        updated_document_url = project.documentUrl
        doc_path = _resolve_docx_path(project.documentUrl or "")
        if doc_path and doc_path.exists():
            doc = Document(doc_path)
            doc.add_page_break()
            doc.add_heading(caption, level=1)
            doc.add_paragraph("Generated from DocuVerse Studio.")
            doc.add_picture(str(image_path), width=Inches(6.5))
            doc.save(doc_path)
        else:
            updated_document_url = project.documentUrl

        ProjectStore.save_project(project)
        return JSONResponse(content={
            "imageUrl": image_url,
            "documentUrl": updated_document_url,
            "contentMarkdown": project.contentMarkdown
        })
    except FileNotFoundError as e:
        raise HTTPException(status_code=500, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Append diagram failed: {e}")

# --- DocuVerse Studio & Project Endpoints ---

from backend.beta.models.project import Project, ReviewFeedback, WorkflowEvent
from backend.beta.services.project_store import ProjectStore
import uuid

class CreateProjectRequest(BaseModel):
    id: Optional[str] = None
    name: str
    content: str
    documentUrl: Optional[str] = None
    status: Optional[str] = None
    reviewedDocumentUrl: Optional[str] = None
    clientEmail: Optional[str] = None
    workflowEvents: Optional[list] = None
    reviewFeedback: Optional[list] = None

class ReviewRequest(BaseModel):
    projectId: str
    clientEmail: str
    documentLink: Optional[str] = None
    senderEmail: Optional[str] = None
    projectName: Optional[str] = None
    insights: Optional[list] = None
    notes: Optional[str] = None
    isUpdate: Optional[bool] = False

class WebhookCallbackRequest(BaseModel):
    projectId: str
    action: str # APPROVED, REJECTED
    feedbackText: Optional[str] = None

@app.post("/api/project/create")
async def create_project(request: CreateProjectRequest):
    project_id = request.id or str(uuid.uuid4())
    project = Project(
        id=project_id,
        name=request.name,
        contentMarkdown=request.content,
        documentUrl=request.documentUrl
    )
    if request.status:
        project.status = request.status
    if request.reviewedDocumentUrl:
        project.reviewedDocumentUrl = request.reviewedDocumentUrl
    if request.clientEmail:
        project.clientEmail = request.clientEmail
    if request.workflowEvents:
        project.workflowEvents = request.workflowEvents
    if request.reviewFeedback:
        project.reviewFeedback = request.reviewFeedback
    ProjectStore.save_project(project)
    return JSONResponse(content={"id": project_id, "message": "Project created successfully"})

@app.get("/api/project/{project_id}")
async def get_project(project_id: str):
    project = ProjectStore.get_project(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    return project

@app.post("/api/workflow/start-review")
async def start_review(request: ReviewRequest):
    print(f"DEBUG: start_review triggered for project: {request.projectId}")
    project = ProjectStore.get_project(request.projectId)
    if not project:
        print(f"ERROR: Project {request.projectId} not found in store!")
        raise HTTPException(status_code=404, detail="Project not found in Python backend store")

    # Update Status
    ProjectStore.update_status(project.id, "IN_REVIEW")
    project.clientEmail = request.clientEmail
    project.reviewToken = project.reviewToken or str(uuid.uuid4())
    project.reviewTokenUsed = False
    project.workflowEvents.append(WorkflowEvent(
        date=datetime.now().isoformat(),
        title="Review Started",
        description="Email sent to client for review.",
        status="IN_REVIEW"
    ))
    
    from backend.beta.utils.email_service import send_review_email

    doc_link = request.documentLink or project.documentUrl or f"/download_srs/{project.id}_SRS.docx"
    project_title = request.projectName or project.name or "DocuVerse Project"
    sender = request.senderEmail
    insights = request.insights or []
    notes = request.notes or ""

    insights_block = ""
    if insights:
        lines = []
        for item in insights:
            title = item.get("title", "Insight")
            desc = item.get("desc", "")
            lines.append(f"- {title}: {desc}")
        insights_block = "\nInsights:\n" + "\n".join(lines)

    update_line = "This is a revised version of the document based on previous feedback." if request.isUpdate else "A new SRS document has been generated and is ready for your technical review."
    body = f"""Hello,

You have been invited to review the Software Requirements Specification (SRS) for: {project_title}

Identifying Tag: [TECHNICAL REVIEW STAGE]

Document link: {doc_link}
{insights_block}

Additonal Notes:
{notes or "No additional notes provided."}

{update_line}

Please use the buttons below in the HTML version of this email to Approve or Request Changes.
"""

    warning = None
    try:
        send_review_email(
            to_email=request.clientEmail,
            subject=f"[Action Required] DocuVerse Review{' Update' if request.isUpdate else ''}: {project_title}",
            body=body,
            reply_to=sender,
            document_link=doc_link,
            project_id=project.id,
            review_token=project.reviewToken,
            is_update=request.isUpdate
        )
    except Exception as e:
        import traceback
        warning = str(e)
        full_error = traceback.format_exc()
        print(f"Email send failed: {warning}")
        print(full_error)
        # return fuller error for debug if needed
        # response["debug_error"] = full_error 

    response = {
        "status": "IN_REVIEW",
        "message": f"Review started for {request.clientEmail}"
    }
    if warning:
        response["warning"] = warning
    return JSONResponse(content=response)

@app.post("/api/workflow/resend-review")
async def resend_review(request: ReviewRequest):
    print(f"DEBUG: resend_review triggered for project: {request.projectId}")
    project = ProjectStore.get_project(request.projectId)
    if not project:
        print(f"ERROR: Project {request.projectId} not found in store!")
        raise HTTPException(status_code=404, detail="Project not found in Python backend store")

    ProjectStore.update_status(project.id, "IN_REVIEW")
    project.clientEmail = request.clientEmail
    project.reviewToken = project.reviewToken or str(uuid.uuid4())
    project.reviewTokenUsed = False
    project.workflowEvents.append(WorkflowEvent(
        date=datetime.now().isoformat(),
        title="Review Resent",
        description="Updated review email sent to client.",
        status="IN_REVIEW"
    ))

    from backend.beta.utils.email_service import send_review_email

    doc_link = request.documentLink or project.documentUrl or f"/download_srs/{project.id}_SRS.docx"
    project_title = request.projectName or project.name or "DocuVerse Project"
    sender = request.senderEmail
    insights = request.insights or []
    notes = request.notes or ""

    insights_block = ""
    if insights:
        lines = []
        for item in insights:
            title = item.get("title", "Insight")
            desc = item.get("desc", "")
            lines.append(f"- {title}: {desc}")
        insights_block = "\nInsights:\n" + "\n".join(lines)

    body = f"""Hello,

An updated review has been requested for the Software Requirements Specification (SRS) for: {project_title}

Identifying Tag: [RE-REVIEW STAGE]

Document link: {doc_link}
{insights_block}

Additonal Notes:
{notes or "No additional notes provided."}

Please reply to this email with your review comments or use the Approve/Request Changes buttons.
"""

    warning = None
    try:
        send_review_email(
            to_email=request.clientEmail,
            subject=f"[Revised] DocuVerse Review Update: {project_title}",
            body=body,
            reply_to=sender,
            document_link=doc_link,
            project_id=project.id,
            review_token=project.reviewToken,
            is_update=True
        )
    except Exception as e:
        warning = str(e)
        print(f"Email resend failed: {warning}")

    response = {
        "status": "IN_REVIEW",
        "message": f"Review resent to {request.clientEmail}"
    }
    if warning:
        response["warning"] = warning
    return JSONResponse(content=response)

@app.post("/api/workflow/webhook-callback")
async def webhook_callback(request: WebhookCallbackRequest):
    project = ProjectStore.get_project(request.projectId)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
        
    if request.action == "APPROVED":
        ProjectStore.update_status(project.id, "APPROVED")
    elif request.action == "REJECTED":
        ProjectStore.update_status(project.id, "CHANGES_REQUESTED")
        if request.feedbackText:
            feedback = ReviewFeedback(
                date=datetime.now().isoformat(),
                comment=request.feedbackText,
                source="Client"
            )
            ProjectStore.add_feedback(project.id, feedback)
            
    return JSONResponse(content={"message": "Callback processed", "new_status": project.status})

@app.get("/api/workflow/review")
async def review_from_email(projectId: str, action: str, token: Optional[str] = None):
    project = ProjectStore.get_project(projectId)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    if project.reviewToken and token != project.reviewToken:
        raise HTTPException(status_code=403, detail="Invalid review token")
    if project.reviewTokenUsed:
        return HTMLResponse(content="<h2>Review already processed</h2><p>This review link has already been used.</p>")

    if action == "APPROVED":
        ProjectStore.update_status(project.id, "APPROVED")
        project.reviewTokenUsed = True
        project.workflowEvents.append(WorkflowEvent(
            date=datetime.now().isoformat(),
            title="Approved",
            description="Client approved the document.",
            status="APPROVED"
        ))
        return HTMLResponse(content="<h2>Review recorded: Approved</h2><p>You can close this tab.</p>")
    if action == "REJECTED":
        ProjectStore.update_status(project.id, "CHANGES_REQUESTED")
        project.workflowEvents.append(WorkflowEvent(
            date=datetime.now().isoformat(),
            title="Changes Requested",
            description="Client requested changes.",
            status="CHANGES_REQUESTED"
        ))
        return HTMLResponse(content=f"""
        <html>
          <body style="font-family: Arial, sans-serif; background:#0f141b; color:#e6edf3; padding:24px;">
            <h2>Changes Requested</h2>
            <p>Please add your review notes below:</p>
            <form method="post" action="/api/workflow/review-feedback">
              <input type="hidden" name="projectId" value="{project.id}" />
              <input type="hidden" name="token" value="{token or ''}" />
              <textarea name="feedbackText" rows="6" style="width:100%; max-width:720px; padding:12px; border-radius:8px; border:1px solid #30363d; background:#161b22; color:#e6edf3;"></textarea>
              <div style="margin-top:12px;">
                <button type="submit" style="background:#dc2626; color:white; border:none; padding:10px 16px; border-radius:6px;">Send Feedback</button>
              </div>
            </form>
          </body>
        </html>
        """)

    raise HTTPException(status_code=400, detail="Invalid action")

@app.post("/api/workflow/review-feedback")
async def review_feedback(projectId: str = Form(...), feedbackText: str = Form(...), token: str = Form("")):
    project = ProjectStore.get_project(projectId)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    if project.reviewToken and token != project.reviewToken:
        raise HTTPException(status_code=403, detail="Invalid review token")
    ProjectStore.update_status(project.id, "CHANGES_REQUESTED")
    feedback = ReviewFeedback(
        date=datetime.now().isoformat(),
        comment=feedbackText,
        source="Client"
    )
    ProjectStore.add_feedback(project.id, feedback)
    project.reviewTokenUsed = True
    project.workflowEvents.append(WorkflowEvent(
        date=datetime.now().isoformat(),
        title="Feedback Received",
        description="Client submitted feedback.",
        status="CHANGES_REQUESTED"
    ))
    return HTMLResponse(content="<h2>Feedback received</h2><p>Thank you. You can close this tab.</p>")














@app.post("/api/project/{project_id}/upload-review")
async def upload_review(project_id: str, file: UploadFile = File(...)):
    project = ProjectStore.get_project(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    # Save file
    filename = f"{project_id}_reviewed_{file.filename}"
    file_path = Path("backend/beta/static") / filename
    with open(file_path, "wb") as buffer:
        import shutil
        shutil.copyfileobj(file.file, buffer)
        
    # Update project
    ProjectStore.update_status(project.id, "CHANGES_REQUESTED")
    project.reviewedDocumentUrl = f"/static/{filename}"
    
    # Add system feedback
    feedback = ReviewFeedback(
        date=datetime.now().isoformat(),
        comment="Client uploaded a marked-up document with changes.",
        source="Client Attachment"
    )
    ProjectStore.add_feedback(project.id, feedback)

    return JSONResponse(content={
        "status": "CHANGES_REQUESTED",
        "reviewedDocumentUrl": project.reviewedDocumentUrl,
        "message": "Review document uploaded successfully"
    })
